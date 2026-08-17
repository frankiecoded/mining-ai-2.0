"""Comprehensive Knowledge Base for Mining AI Platform.

Indexes, stores, and retrieves all documents shared with the platform.
Pure Python implementation with no heavy dependencies.
"""

import json
import os
import hashlib
import logging
import re
import math
from datetime import datetime, timezone
from typing import Any, Optional
from dataclasses import dataclass, field, asdict

logger = logging.getLogger(__name__)

DATA_DIR = "/opt/aios/data"
KB_PATH = os.path.join(DATA_DIR, "knowledge_base.json")

MINING_VOCABULARY = {
    "mining", "ore", "mineral", "deposit", "grade", "tonnage", "metallurgy",
    "comminution", "flotation", "leaching", "heap", "tailings", "slurry",
    "drill", "blast", "excavation", "haulage", "overburden", "bench",
    "pit", "underground", "shaft", "ramp", "stope", "caving",
    "geotechnical", "rock", "mass", "fault", "fracture", "joint",
    "dip", "strike", "lithology", "stratigraphy", "sedimentary",
    "igneous", "metamorphic", " alteration", "mineralization",
    "assay", "ppm", "g/t", "品位", "recovery", "processing",
    "crushing", "grinding", "mill", "autogenous", "ball", "rod",
    "concentrate", "smelting", "refining", "electrowinning",
    "water", "dust", "emission", "rehabilitation", "closure",
    "feasibility", "pre-feasibility", "bankable", "reserves",
    "resources", "measured", "indicated", "inferred",
    "JORC", "NI 43-101", "SAMREC", "compliance", "permit",
    "environmental", "impact", "assessment", "EIA",
    "satellite", "remote", "sensing", "multispectral", "hyperspectral",
    "lidar", "DEM", "orthophoto", "NDVI", "InSAR",
    "blast", "vibration", "seismic", "ground", "control",
    "haul", "truck", "loader", "conveyor", "fleet",
    "safety", "incident", "hazard", "risk", "emergency",
    "cost", "capex", "opex", "NPV", "IRR", "payback",
    "geology", "geophysical", "geochemical", "geospatial",
    "GIS", "coordinate", "UTM", "latitude", "longitude",
    "geotechnical", "slope", "stability", "displacement", "deformation",
    "spectral", "band", "pixel", "resolution", "acquisition",
    "lithology", " alteration", "vein", "reef", "lode",
    "placer", "alluvial", "colluvial", "saprolite",
    "core", "HQ", "NQ", "diamond", "percussion",
    "survey", "topographic", "contour", "elevation",
    "block", "model", "kriging", "interpolation", "estimation",
    "mine", "plan", "design", "sequence", "production",
    "stockpile", " ROM", "crusher", "sizer", "dumper",
    "pipeline", "infrastructure", "power", "water", "access",
    "community", "stakeholder", "indigenous", "land",
    "insurance", "financial", "model", "valuation",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".geotiff"}
MIME_MAP = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".geotiff": "image/tiff",
    ".json": "application/json",
}


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _hash_file(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()[:16]


def _compute_relevance(query_tokens: list[str], text_tokens: list[str]) -> float:
    if not query_tokens or not text_tokens:
        return 0.0
    text_len = len(text_tokens)
    tf_map: dict[str, float] = {}
    for t in text_tokens:
        tf_map[t] = tf_map.get(t, 0.0) + 1.0
    for k in tf_map:
        tf_map[k] = 1.0 + math.log(tf_map[k]) if tf_map[k] > 0 else 0.0
    max_tf = max(tf_map.values()) if tf_map else 1.0
    score = 0.0
    matched = 0
    for qt in query_tokens:
        if qt in tf_map:
            matched += 1
            score += tf_map[qt] / max_tf
    if matched == 0:
        return 0.0
    coverage = matched / len(query_tokens)
    avg_tf = score / matched
    return round(coverage * 0.6 + avg_tf * 0.4, 4)


@dataclass
class KnowledgeDocument:
    doc_id: str = ""
    filename: str = ""
    original_filename: str = ""
    file_type: str = ""
    file_size: int = 0
    mime_type: str = ""

    title: str = ""
    description: str = ""
    tags: list[str] = field(default_factory=list)
    category: str = "other"

    content_text: str = ""
    content_summary: str = ""

    metadata: dict[str, Any] = field(default_factory=dict)

    source: str = "upload"
    session_id: str = ""
    referenced_in: list[str] = field(default_factory=list)
    thumbnail_base64: str = ""

    indexing_status: str = "pending"
    created_at: str = ""
    updated_at: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "KnowledgeDocument":
        known = {f.name for f in cls.__dataclass_fields__.values()}
        filtered = {k: v for k, v in data.items() if k in known}
        return cls(**filtered)


class KnowledgeBase:
    def __init__(self, data_dir: str = DATA_DIR, kb_path: str = KB_PATH):
        self.data_dir = data_dir
        self.kb_path = kb_path
        self.documents: dict[str, KnowledgeDocument] = {}
        os.makedirs(self.data_dir, exist_ok=True)
        self.load()

    def add_document(
        self,
        file_path: str,
        filename: str,
        metadata: Optional[dict[str, Any]] = None,
        session_id: str = "",
        source: str = "upload",
    ) -> KnowledgeDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(filename)[1].lower()
        doc_id = _hash_file(file_path) + ext.replace(".", "")
        file_size = os.path.getsize(file_path)
        mime_type = MIME_MAP.get(ext, "application/octet-stream")
        now = _now_iso()

        doc = KnowledgeDocument(
            doc_id=doc_id,
            filename=os.path.basename(file_path),
            original_filename=filename,
            file_type=ext.lstrip("."),
            file_size=file_size,
            mime_type=mime_type,
            source=source,
            session_id=session_id,
            metadata=metadata or {},
            created_at=now,
            updated_at=now,
        )

        try:
            if ext == ".pdf":
                doc.content_text = self.extract_text_from_pdf(file_path)
            elif ext == ".docx":
                doc.content_text = self.extract_text_from_docx(file_path)
            elif ext == ".xlsx":
                doc.content_text = self.extract_text_from_xlsx(file_path)
            elif ext == ".csv":
                doc.content_text = self.extract_text_from_csv(file_path)
            elif ext == ".txt":
                doc.content_text = self._extract_text_from_txt(file_path)
            elif ext in IMAGE_EXTENSIONS:
                img_meta = self.extract_text_from_image(file_path)
                doc.metadata.update(img_meta)
                doc.content_text = f"Image: {filename}. {json.dumps(img_meta)}"
            elif ext == ".json":
                doc.content_text = self._extract_text_from_json(file_path)
            else:
                doc.content_text = f"Unsupported file type: {ext}"
                doc.indexing_status = "failed"
                self.documents[doc_id] = doc
                self.save()
                return doc

            doc.title = self._derive_title(filename, doc.content_text)
            doc.content_summary = self.generate_summary(doc.content_text)
            doc.tags = self.extract_key_terms(doc.content_text)
            doc.category = self._auto_categorize(doc.content_text, ext)
            doc.indexing_status = "indexed"

            if "page_count" not in doc.metadata and ext == ".pdf":
                doc.metadata["page_count"] = doc.content_text.count("\f") + 1
            doc.metadata["word_count"] = len(doc.content_text.split())
            if metadata:
                doc.metadata.update(metadata)

        except Exception as exc:
            logger.error("Failed to index %s: %s", filename, exc)
            doc.indexing_status = "failed"
            doc.content_summary = f"Indexing failed: {exc}"

        self.documents[doc_id] = doc
        self.save()
        logger.info("Added document %s (%s) as %s", filename, doc.file_type, doc_id)
        return doc

    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            pages: list[str] = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append(f"--- Page {i + 1} ---\n{text}")
            return "\n\n".join(pages)
        except ImportError:
            return self._extract_text_from_pdf_fallback(file_path)
        except Exception as exc:
            logger.warning("pypdf failed for %s: %s, trying fallback", file_path, exc)
            return self._extract_text_from_pdf_fallback(file_path)

    def _extract_text_from_pdf_fallback(self, file_path: str) -> str:
        try:
            with open(file_path, "rb") as f:
                raw = f.read()
            text_parts: list[str] = []
            current: list[bytes] = []
            for byte in raw:
                current.append(bytes([byte]))
                if byte == 0x0A:
                    try:
                        decoded = b"".join(current).decode("latin-1", errors="ignore")
                        if any(c.isalpha() for c in decoded):
                            text_parts.append(decoded.strip())
                    except Exception:
                        pass
                    current = []
            joined = " ".join(text_parts)
            cleaned = re.sub(r"\s+", " ", joined).strip()
            if len(cleaned) > 100:
                return cleaned[:50000]
            return f"[PDF binary - text extraction unavailable. File: {os.path.basename(file_path)}]"
        except Exception as exc:
            return f"[PDF extraction failed: {exc}]"

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    tables_text.append(" | ".join(cells))
            parts = paragraphs + (["--- Tables ---"] + tables_text if tables_text else [])
            return "\n".join(parts)
        except ImportError:
            return self._extract_text_from_docx_fallback(file_path)
        except Exception as exc:
            logger.warning("python-docx failed for %s: %s", file_path, exc)
            return self._extract_text_from_docx_fallback(file_path)

    def _extract_text_from_docx_fallback(self, file_path: str) -> str:
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            with zipfile.ZipFile(file_path, "r") as z:
                with z.open("word/document.xml") as f:
                    tree = ET.parse(f)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            texts: list[str] = []
            for node in tree.iter(f"{{{ns['w']}}}t"):
                if node.text:
                    texts.append(node.text)
            return " ".join(texts) if texts else "[DOCX: no text extracted]"
        except Exception as exc:
            return f"[DOCX extraction failed: {exc}]"

    def extract_text_from_xlsx(self, file_path: str) -> str:
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            with zipfile.ZipFile(file_path, "r") as z:
                with z.open("xl/workbook.xml") as f:
                    wb_tree = ET.parse(f)
            ns_main = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
            sheet_names: list[str] = []
            for sname_el in wb_tree.iter(
                f"{{{ns_main['main']}}}sheet"
            ):
                name_attr = sname_el.get("name", "Sheet")
                sheet_names.append(name_attr)

            parts: list[str] = [f"Sheets: {', '.join(sheet_names)}"]
            try:
                with zipfile.ZipFile(file_path, "r") as z:
                    with z.open("xl/sharedStrings.xml") as f:
                        ss_tree = ET.parse(f)
                strings: list[str] = []
                for si in ss_tree.iter(f"{{{ns_main['main']}}}t"):
                    if si.text:
                        strings.append(si.text)
                if strings:
                    parts.append(f"Shared strings count: {len(strings)}")
                    sample = strings[:200]
                    parts.append("Sample data: " + " | ".join(sample[:50]))
            except (KeyError, FileNotFoundError):
                parts.append("[No shared strings found]")

            try:
                with zipfile.ZipFile(file_path, "r") as z:
                    with z.open("xl/worksheets/sheet1.xml") as f:
                        ws_tree = ET.parse(f)
                row_count = sum(1 for _ in ws_tree.iter(f"{{{ns_main['main']}}}row"))
                parts.append(f"Row count (sheet1): {row_count}")
            except (KeyError, FileNotFoundError):
                pass

            return "\n".join(parts)
        except Exception as exc:
            return f"[XLSX extraction failed: {exc}]"

    def extract_text_from_csv(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                lines = []
                for i, line in enumerate(f):
                    lines.append(line.rstrip("\n\r"))
                    if i >= 500:
                        break
            if not lines:
                return "[Empty CSV file]"

            headers = lines[0]
            row_count = sum(1 for _ in open(file_path, "r", encoding="utf-8", errors="replace")) - 1
            sample_rows = lines[1:6]

            parts = [
                f"CSV Headers: {headers}",
                f"Row count: {row_count}",
                f"Column count: {len(headers.split(','))}",
                "--- Sample data (first 5 rows) ---",
            ]
            for row in sample_rows:
                parts.append(row)
            return "\n".join(parts)
        except Exception as exc:
            return f"[CSV extraction failed: {exc}]"

    def extract_text_from_image(self, file_path: str) -> dict[str, Any]:
        info: dict[str, Any] = {
            "format": os.path.splitext(file_path)[1].lstrip(".").upper(),
            "is_satellite": False,
            "is_imagery": True,
        }
        try:
            with open(file_path, "rb") as f:
                header = f.read(32)
            if len(header) >= 8:
                if header[:8] == b"\x89PNG\r\n\x1a\n":
                    info["format"] = "PNG"
                    info["width"] = int.from_bytes(header[16:20], "big")
                    info["height"] = int.from_bytes(header[20:24], "big")
                    info["bit_depth"] = header[24]
                    info["color_type"] = header[25]
                    channels = {0: 1, 2: 3, 3: 1, 4: 2, 6: 4}
                    info["bands"] = channels.get(info.get("color_type", 0), 3)
                elif header[:2] == b"\xff\xd8":
                    info["format"] = "JPEG"
                elif header[:4] == b"II\x2a\x00" or header[:4] == b"MM\x00\x2a":
                    info["format"] = "TIFF"
                    info["is_satellite"] = True
                    info["bands"] = 4
                else:
                    info["format"] = "UNKNOWN"

            file_size = os.path.getsize(file_path)
            info["file_size"] = file_size

            filename_lower = os.path.basename(file_path).lower()
            satellite_keywords = [
                "sentinel", "landsat", "aster", "modis", "spot",
                "worldview", "geotiff", "ortho", "ndvi", "aerial",
                "satellite", "imagery", "rgb", "multispectral",
            ]
            if any(kw in filename_lower for kw in satellite_keywords):
                info["is_satellite"] = True
            if info["format"] in ("TIFF",) and info.get("bands", 1) >= 3:
                info["is_satellite"] = True

        except Exception as exc:
            info["error"] = str(exc)

        return info

    def _extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception as exc:
            return f"[TXT extraction failed: {exc}]"

    def _extract_text_from_json(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)

            parts: list[str] = ["JSON Structure:"]
            if isinstance(data, dict):
                parts.append(f"Top-level keys: {list(data.keys())}")
                for key in list(data.keys())[:10]:
                    val = data[key]
                    vtype = type(val).__name__
                    if isinstance(val, (list, dict)):
                        vlen = len(val)
                        parts.append(f"  {key}: {vtype} (length={vlen})")
                    else:
                        parts.append(f"  {key}: {vtype} = {str(val)[:100]}")
            elif isinstance(data, list):
                parts.append(f"Array of {len(data)} items")
                if data:
                    parts.append(f"Item type: {type(data[0]).__name__}")
                    parts.append(f"Sample: {str(data[0])[:200]}")
            else:
                parts.append(f"Value: {str(data)[:500]}")
            return "\n".join(parts)
        except Exception as exc:
            return f"[JSON extraction failed: {exc}]"

    def generate_summary(self, text: str) -> str:
        if not text:
            return "[No content]"
        words = text.split()
        word_count = len(words)
        truncated = " ".join(words[:500])
        key_terms = self.extract_key_terms(text)
        summary = f"Word count: {word_count}. Key terms: {', '.join(key_terms[:15])}.\n\n{truncated}"
        if len(summary) > 2000:
            summary = summary[:2000] + "..."
        return summary

    def extract_key_terms(self, text: str) -> list[str]:
        if not text:
            return []
        words = re.findall(r"[a-zA-Z]{3,}", text.lower())
        freq: dict[str, int] = {}
        for w in words:
            freq[w] = freq.get(w, 0) + 1

        mining_hits: dict[str, int] = {}
        for term in MINING_VOCABULARY:
            term_lower = term.lower().strip()
            if term_lower in freq:
                mining_hits[term_lower] = freq[term_lower]

        non_stopwords = {
            k: v for k, v in freq.items()
            if v >= 2 and k not in _STOP_WORDS and len(k) >= 4
        }
        general_top = sorted(non_stopwords.items(), key=lambda x: -x[1])[:20]

        combined: dict[str, int] = {}
        for term, count in mining_hits.items():
            combined[term] = count * 3
        for term, count in general_top:
            combined[term] = combined.get(term, 0) + count

        ranked = sorted(combined.items(), key=lambda x: -x[1])
        return [t for t, _ in ranked[:20]]

    def _auto_categorize(self, text: str, file_type: str) -> str:
        text_lower = text.lower()

        if file_type in IMAGE_EXTENSIONS:
            if any(kw in text_lower for kw in ["satellite", "sentinel", "landsat", "ndvi", "imagery"]):
                return "satellite"
            return "satellite"

        category_keywords: dict[str, list[str]] = {
            "geological": [
                "geology", "lithology", "mineral", "ore", "deposit", "grade",
                "geophysical", "geochemical", "assay", "core", "drill",
                "stratigraphy", "fault", "vein", "reef", "alteration",
                "mineralization", "geospatial", "geographic",
            ],
            "geotechnical": [
                "geotechnical", "slope", "stability", "rock mass", "fault",
                "fracture", "joint", "displacement", "deformation", "ground control",
                "seismic", "vibration", "blast", "excavation",
            ],
            "operational": [
                "production", "mine plan", "haul", "truck", "fleet",
                "processing", "crushing", "grinding", "mill", "conveyor",
                "stockpile", "infrastructure", "safety", "incident",
            ],
            "environmental": [
                "environmental", "water", "dust", "emission", "rehabilitation",
                "closure", "tailings", "waste", "conservation",
                "impact assessment", "biodiversity",
            ],
            "financial": [
                "financial", "cost", "capex", "opex", "npv", "irr",
                "feasibility", "budget", "revenue", "valuation",
                "payback", "profit", "economics",
            ],
            "regulatory": [
                "compliance", "permit", "regulation", "jorc", "ni 43-101",
                "samrec", "audit", "inspection", "license", "law",
                "environmental impact", "eia",
            ],
            "satellite": [
                "satellite", "remote sensing", "multispectral", "hyperspectral",
                "lidar", "dem", "orthophoto", "ndvi", "insar",
                "sentinel", "landsat", "imagery", "pixel", "resolution",
            ],
        }

        scores: dict[str, float] = {}
        for cat, keywords in category_keywords.items():
            score = sum(text_lower.count(kw) for kw in keywords)
            scores[cat] = score

        if scores:
            best = max(scores, key=lambda c: scores[c])
            if scores[best] > 0:
                return best
        return "other"

    def _derive_title(self, filename: str, content_text: str) -> str:
        base = os.path.splitext(filename)[0]
        cleaned = re.sub(r"[_\-]+", " ", base).strip()
        cleaned = re.sub(r"\s+", " ", cleaned)
        if cleaned and len(cleaned) > 3:
            return cleaned.title()
        first_line = content_text.split("\n")[0][:100].strip()
        if first_line:
            return first_line
        return filename

    def search(self, query: str) -> list[dict[str, Any]]:
        if not query.strip():
            return []
        query_tokens = [t.lower() for t in re.findall(r"[a-zA-Z0-9]{2,}", query)]
        if not query_tokens:
            return []

        results: list[dict[str, Any]] = []
        for doc in self.documents.values():
            searchable = " ".join([
                doc.original_filename,
                doc.title,
                doc.description,
                doc.content_text[:10000],
                " ".join(doc.tags),
                doc.category,
            ]).lower()
            text_tokens = [t for t in re.findall(r"[a-zA-Z0-9]{2,}", searchable)]

            score = _compute_relevance(query_tokens, text_tokens)

            if score > 0:
                if any(qt in doc.original_filename.lower() for qt in query_tokens):
                    score += 0.3
                if any(qt in doc.title.lower() for qt in query_tokens):
                    score += 0.2
                if any(qt in " ".join(doc.tags).lower() for qt in query_tokens):
                    score += 0.15

                results.append({
                    "doc_id": doc.doc_id,
                    "filename": doc.original_filename,
                    "title": doc.title,
                    "category": doc.category,
                    "file_type": doc.file_type,
                    "relevance_score": round(min(score, 1.0), 4),
                    "snippet": doc.content_summary[:200],
                })

        results.sort(key=lambda r: -r["relevance_score"])
        return results

    def search_by_category(self, category: str) -> list[dict[str, Any]]:
        cat_lower = category.lower()
        return [
            {
                "doc_id": d.doc_id,
                "filename": d.original_filename,
                "title": d.title,
                "file_type": d.file_type,
                "created_at": d.created_at,
                "file_size": d.file_size,
            }
            for d in self.documents.values()
            if d.category.lower() == cat_lower
        ]

    def search_by_type(self, file_type: str) -> list[dict[str, Any]]:
        ft = file_type.lower().lstrip(".")
        return [
            {
                "doc_id": d.doc_id,
                "filename": d.original_filename,
                "title": d.title,
                "category": d.category,
                "created_at": d.created_at,
                "file_size": d.file_size,
            }
            for d in self.documents.values()
            if d.file_type.lower() == ft
        ]

    def search_by_session(self, session_id: str) -> list[dict[str, Any]]:
        return [
            {
                "doc_id": d.doc_id,
                "filename": d.original_filename,
                "title": d.title,
                "category": d.category,
                "file_type": d.file_type,
                "created_at": d.created_at,
            }
            for d in self.documents.values()
            if d.session_id == session_id
        ]

    def get_document(self, doc_id: str) -> Optional[KnowledgeDocument]:
        return self.documents.get(doc_id)

    def list_all_documents(self) -> list[dict[str, Any]]:
        docs = sorted(
            self.documents.values(),
            key=lambda d: d.created_at,
            reverse=True,
        )
        return [
            {
                "doc_id": d.doc_id,
                "filename": d.original_filename,
                "file_type": d.file_type,
                "category": d.category,
                "title": d.title,
                "created_at": d.created_at,
                "file_size": d.file_size,
                "indexing_status": d.indexing_status,
            }
            for d in docs
        ]

    def get_statistics(self) -> dict[str, Any]:
        total_docs = len(self.documents)
        total_size = sum(d.file_size for d in self.documents.values())

        type_counts: dict[str, int] = {}
        category_counts: dict[str, int] = {}
        for d in self.documents.values():
            type_counts[d.file_type] = type_counts.get(d.file_type, 0) + 1
            category_counts[d.category] = category_counts.get(d.category, 0) + 1

        recent = sorted(
            self.documents.values(),
            key=lambda d: d.created_at,
            reverse=True,
        )[:10]

        return {
            "total_documents": total_docs,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "documents_by_type": type_counts,
            "documents_by_category": category_counts,
            "recent_uploads": [
                {
                    "doc_id": d.doc_id,
                    "filename": d.original_filename,
                    "category": d.category,
                    "created_at": d.created_at,
                }
                for d in recent
            ],
        }

    def get_recent_documents(self, limit: int = 10) -> list[dict[str, Any]]:
        recent = sorted(
            self.documents.values(),
            key=lambda d: d.created_at,
            reverse=True,
        )[:limit]
        return [
            {
                "doc_id": d.doc_id,
                "filename": d.original_filename,
                "title": d.title,
                "file_type": d.file_type,
                "category": d.category,
                "created_at": d.created_at,
                "file_size": d.file_size,
                "description": d.description[:200] if d.description else "",
            }
            for d in recent
        ]

    def tag_document(self, doc_id: str, tags: list[str]) -> bool:
        doc = self.documents.get(doc_id)
        if not doc:
            return False
        existing = set(doc.tags)
        for tag in tags:
            normalized = tag.strip().lower()
            if normalized and normalized not in existing:
                doc.tags.append(normalized)
        doc.updated_at = _now_iso()
        self.save()
        return True

    def update_document(self, doc_id: str, updates: dict[str, Any]) -> bool:
        doc = self.documents.get(doc_id)
        if not doc:
            return False
        allowed_fields = {
            "title", "description", "tags", "category",
            "source", "session_id", "metadata",
        }
        for key, value in updates.items():
            if key in allowed_fields:
                setattr(doc, key, value)
        doc.updated_at = _now_iso()
        self.save()
        return True

    def delete_document(self, doc_id: str) -> bool:
        if doc_id in self.documents:
            del self.documents[doc_id]
            self.save()
            logger.info("Deleted document %s", doc_id)
            return True
        return False

    def reference_document(self, doc_id: str, message_id: str) -> bool:
        doc = self.documents.get(doc_id)
        if not doc:
            return False
        if message_id not in doc.referenced_in:
            doc.referenced_in.append(message_id)
            doc.updated_at = _now_iso()
            self.save()
        return True

    def get_referenced_documents(self, message_ids: list[str]) -> list[dict[str, Any]]:
        msg_set = set(message_ids)
        results: list[dict[str, Any]] = []
        for doc in self.documents.values():
            refs = msg_set.intersection(doc.referenced_in)
            if refs:
                results.append({
                    "doc_id": doc.doc_id,
                    "filename": doc.original_filename,
                    "title": doc.title,
                    "referenced_by_messages": list(refs),
                })
        return results

    def get_knowledge_summary(self) -> dict[str, Any]:
        stats = self.get_statistics()
        total = stats["total_documents"]

        recent_activity: list[dict[str, Any]] = []
        now = datetime.now(timezone.utc)
        for doc in self.documents.values():
            try:
                doc_time = datetime.fromisoformat(doc.created_at)
                days_ago = (now - doc_time).days
                if days_ago <= 7:
                    recent_activity.append({
                        "doc_id": doc.doc_id,
                        "filename": doc.original_filename,
                        "days_ago": days_ago,
                    })
            except (ValueError, TypeError):
                pass

        all_terms: dict[str, int] = {}
        for doc in self.documents.values():
            for term in doc.tags[:10]:
                all_terms[term] = all_terms.get(term, 0) + 1
        top_topics = sorted(all_terms.items(), key=lambda x: -x[1])[:20]

        categories = stats["documents_by_category"]
        most_active = max(categories, key=lambda c: categories[c]) if categories else "none"

        return {
            "total_documents": total,
            "total_size_mb": stats["total_size_mb"],
            "categories": categories,
            "types": stats["documents_by_type"],
            "most_active_category": most_active,
            "top_topics": [{"term": t, "frequency": f} for t, f in top_topics],
            "recent_activity_count": len(recent_activity),
            "recent_documents": recent_activity[:10],
            "indexing_health": {
                "indexed": sum(
                    1 for d in self.documents.values() if d.indexing_status == "indexed"
                ),
                "pending": sum(
                    1 for d in self.documents.values() if d.indexing_status == "pending"
                ),
                "failed": sum(
                    1 for d in self.documents.values() if d.indexing_status == "failed"
                ),
            },
        }

    def save(self) -> None:
        try:
            data = {
                "version": "1.0",
                "saved_at": _now_iso(),
                "documents": {
                    doc_id: doc.to_dict()
                    for doc_id, doc in self.documents.items()
                },
            }
            tmp_path = self.kb_path + ".tmp"
            with open(tmp_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            os.replace(tmp_path, self.kb_path)
        except Exception as exc:
            logger.error("Failed to save knowledge base: %s", exc)

    def load(self) -> None:
        if not os.path.exists(self.kb_path):
            logger.info("No existing knowledge base found at %s", self.kb_path)
            return
        try:
            with open(self.kb_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            docs_raw = data.get("documents", {})
            for doc_id, doc_data in docs_raw.items():
                doc = KnowledgeDocument.from_dict(doc_data)
                if not doc.doc_id:
                    doc.doc_id = doc_id
                self.documents[doc_id] = doc
            logger.info(
                "Loaded %d documents from knowledge base", len(self.documents)
            )
        except Exception as exc:
            logger.error("Failed to load knowledge base: %s", exc)
            self.documents = {}


_STOP_WORDS = {
    "the", "and", "for", "are", "but", "not", "you", "all", "can",
    "had", "her", "was", "one", "our", "out", "has", "his", "how",
    "its", "may", "new", "now", "old", "see", "way", "who", "did",
    "get", "let", "say", "she", "too", "use", "that", "with",
    "this", "will", "each", "make", "like", "long", "look",
    "many", "some", "them", "than", "then", "what", "when",
    "your", "said", "there", "been", "have", "from", "they",
    "were", "being", "would", "could", "should", "about",
    "other", "into", "over", "such", "those", "after",
    "also", "just", "only", "very", "well", "back",
    "were", "does", "done", "much", "more", "most",
    "file", "data", "text", "type", "name", "size",
    "document", "content", "based", "using", "used",
    "sheet", "page", "rows", "columns", "format",
}


__all__ = ["KnowledgeBase", "KnowledgeDocument"]
