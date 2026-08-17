"""
Intelligent Document Reading and Understanding Engine for Mining AI.

Provides comprehensive document parsing, entity extraction, content analysis,
and cross-reference capabilities for mining industry documents.
"""

import json
import os
import re
import hashlib
import logging
import struct
import io
import math
from datetime import datetime
from typing import Any, Dict, List, Optional, Set, Tuple
from dataclasses import dataclass, field
from enum import Enum

logger = logging.getLogger(__name__)

MINING_VOCABULARY: Dict[str, List[str]] = {
    "minerals": [
        "gold", "silver", "copper", "iron", "iron ore", "lithium", "cobalt",
        "platinum", "palladium", "uranium", "rare earth", "rare earth elements",
        "diamond", "diamonds", "coal", "coal seam", "bauxite", "nickel", "zinc",
        "lead", "tin", "tungsten", "molybdenum", "chromium", "manganese",
        "titanium", "vanadium", "magnesium", "phosphate", "potash", "sulfur",
        "gypsum", "limestone", "dolomite", "feldspar", "quartz", "mica",
        "kaolin", "bentonite", "garnet", "ilmenite", "zircon", "monazite",
        "chrysotile", "asbestos", "magnetite", "hematite", "goethite",
        "chalcopyrite", "bornite", "galena", "sphalerite", "pyrite", "pyrrhotite",
        "cinnabar", "stibnite", "arsenopyrite", "pentlandite", "chromite",
        "wolframite", "scheelite", "columbite", "tantalite", "beryl",
        "spodumene", "lepidolite", "petalite", "amblygonite",
    ],
    "equipment": [
        "excavator", "haul truck", "dump truck", "loader", "backhoe",
        "drill rig", "drill", "blast hole drill", "rotary drill",
        "crusher", "jaw crusher", "cone crusher", "impact crusher",
        "conveyor", "conveyor belt", "stacker", "reclaimer",
        "ball mill", "sag mill", "rod mill", "grinding mill",
        "flotation cell", "thickener", "filter press",
        "kiln", "rotary kiln", "furnace", "smelter", "refinery",
        "pump", "slurry pump", "centrifugal pump",
        "dozer", "bulldozer", "grader", "scraper", "compactor",
        "auger", "boring machine", "tunnel boring machine",
        "dredge", "dragline", "power shovel", "rope shovel",
        "front-end loader", "wheel loader", "skid steer",
        "water truck", "service truck", "fuel truck",
        "crane", "mobile crane", "overhead crane",
        "separator", "magnetic separator", "gravity separator",
        "leach pad", "heap leach", "vat leach",
        "heap leach pad", "carbon in pulp", "cip",
        "carbon in leach", "cil", "carbon in column",
        "agglomerator", "autoclave", "roaster",
    ],
    "processes": [
        "mining", "excavation", "drilling", "blasting", "hauling",
        "crushing", "grinding", "milling", "comminution",
        "flotation", "froth flotation", "dense media separation",
        "gravity separation", "magnetic separation", "electrostatic separation",
        "leaching", "heap leaching", "in-situ leaching", "cyanidation",
        "ammonia leaching", "acid leaching", "bioleaching",
        "smelting", "electrowinning", "electrorefining", "refining",
        "cyanide", "mercury amalgamation", "carbon adsorption",
        "tailings", "tailings dam", "tailings storage",
        "reclamation", "remediation", "rehabilitation",
        "dewatering", "dewatering system", "water treatment",
        "dust suppression", "dust control", "ventilation",
        "ground support", "slope stability", "rock mechanics",
        "geotechnical", "seismic survey", "geophysical survey",
        "core drilling", "diamond drilling", "reverse circulation",
        "air core drilling", "underground mining", "open pit",
        "surface mining", "sublevel caving", "stoping",
        "room and pillar", "longwall mining", "block caving",
        "cut and fill", "shrinkage stoping", "top slicing",
        "artisanal mining", "small-scale mining",
    ],
    "chemicals": [
        "sodium cyanide", "potassium cyanide", "calcium cyanide",
        "sulfuric acid", "hydrochloric acid", "nitric acid",
        "hydrofluoric acid", "sodium hydroxide", "lime", "quicklime",
        "calcium oxide", "calcium hydroxide", "soda ash",
        "sodium carbonate", "xanthate", "potassium amyl xanthate",
        "sodium isopropyl xanthate", "frother", "mibc",
        "methyl isobutyl carbinol", "collector", "depressant",
        "activator", "flocculant", "polymer", "polyacrylamide",
        "mercury", "thiourea", "thiosulfate",
        "sodium sulfide", "sodium hydrosulfide",
        "copper sulfate", "ferric chloride", "sodium chloride",
        "hydrogen peroxide", "ozone", "chlorine",
        "emulsifier", "surfactant", "wetting agent",
        "diesel", "fuel oil", "kerosene",
    ],
    "regulations": [
        "EPA", "environmental protection", "environmental impact assessment",
        "EIA", "NEPA", "clean air act", "clean water act",
        "OSHA", "occupational safety", "mine safety",
        "MSHA", "mine safety and health administration",
        "permit", "mining permit", "environmental permit",
        "license", "mining license", "exploration license",
        "royalty", "mining royalty", "mineral rights",
        "land use", "land access", "easement",
        "indigenous rights", "native title", "free prior informed consent",
        "FPIC", "community consultation", "stakeholder engagement",
        "tailings management", "water discharge", "air quality",
        "noise ordinance", "vibration limit", "subsidence",
        "waste management", "hazardous waste", "solid waste",
        "mine closure", "decommissioning", "bond",
        "reclamation bond", "closure plan", "post-closure",
        "ISO 14001", "ISO 45001", "EMESRT",
        "ICMM", "responsible mining", "sustainable mining",
        "global reporting initiative", "GRI",
    ],
    "safety_terms": [
        "safety", "hazard", "risk assessment", "risk management",
        "incident", "accident", "near miss", "lost time injury",
        "LTI", "total recordable injury", "TRI",
        "fatality", "serious injury", "first aid",
        "PPE", "personal protective equipment",
        "hard hat", "safety boots", "high visibility",
        "safety glasses", "respirator", "hearing protection",
        "fall protection", "guarding", "lockout tagout", "LOTO",
        "confined space", "atmospheric monitoring",
        "gas detection", "methane", "carbon monoxide", "hydrogen sulfide",
        "fire prevention", "fire suppression", "explosion prevention",
        "ground control", "roof support", "pillar", "span",
        "scaling", "barring down", "rock bolt", "shotcrete",
        "safety meeting", "toolbox talk", "pre-shift inspection",
        "emergency response", "evacuation", "first aid",
        "mine rescue", "self-rescuer", "MSHA training",
        "competent person", "qualified person", "authorized person",
    ],
    "geological_terms": [
        "geology", "geologist", "stratigraphy", "stratigraphic",
        "lithology", "lithological", "structure", "structural",
        "fault", "faulting", "fracture", "joint", "fold",
        "anticline", "syncline", "monocline",
        "ore body", "orebody", "ore zone", "mineralization",
        "grade", "tonnage", "cut-off grade", "average grade",
        "dilution", "recovery", "metallurgical recovery",
        "reserve", "resource", "proved", "probable", "possible",
        "indicated", "inferred", "measured",
        "JORC", "NI 43-101", "CIM", "SAMREC",
        "feasibility study", "pre-feasibility", "scoping study",
        "bankable feasibility", "BFS",
        "drill core", "drill hole", "assay", "assay result",
        "geophysical", "geophysics", "magnetic", "gravity",
        "electromagnetic", "IP", "induced polarization",
        "seismic", "seismic reflection",
        "alteration", "hydrothermal", "magmatic", "sedimentary",
        "metamorphic", "igneous", "intrusive", "extrusive",
        "batholith", "stock", "dike", "sill", "vein",
        "breccia", "skarn", "porphyry", "epithermal",
        "VMS", "volcanogenic massive sulfide",
        "BIF", "banded iron formation",
        "placer", "alluvial", "eluvial", "residual",
        "saprolite", "laterite", "oxide", "sulfide", "mixed",
        "weathering", "supergene", "hypogene",
        "basemap", "cross section", "long section", "plan view",
        "contour", "elevation", "topography", "bathymetry",
    ],
}

DOC_TYPE_PATTERNS: Dict[str, List[str]] = {
    "feasibility_study": [
        r"feasibility\s+study", r"bankable\s+feasibility",
        r"pre[\-\s]feasibility", r"economics?\s+of\s+mining",
        r"capital\s+cost", r"operating\s+cost", r"npv", r"irr",
        r"net\s+present\s+value", r"internal\s+rate\s+of\s+return",
        r"payback\s+period", r"cash\s+flow", r"economic\s+evaluation",
    ],
    "environmental_impact": [
        r"environmental\s+impact", r"EIA", r"environmental\s+assessment",
        r"environmental\s+baseline", r"biodiversity", r"ecosystem",
        r"flora", r"fauna", r"water\s+quality", r"air\s+quality",
        r"noise\s+assessment", r"visual\s+impact", r"heritage",
        r"social\s+impact", r"community\s+impact",
    ],
    "technical_report": [
        r"technical\s+report", r"geological\s+report",
        r"mining\s+report", r"geotechnical\s+report",
        r"metallurgical\s+report", r"exploration\s+report",
        r"drill\s+results", r"resource\s+estimate",
        r"reserve\s+estimate", r"mine\s+plan",
    ],
    "financial_analysis": [
        r"financial\s+analysis", r"financial\s+model",
        r"cost\s+analysis", r"budget", r"forecast",
        r"revenue\s+projection", r"profitability",
        r"income\s+statement", r"balance\s+sheet",
        r"cash\s+flow\s+statement", r"financial\s+statement",
    ],
    "safety_report": [
        r"safety\s+report", r"incident\s+report",
        r"accident\s+report", r"risk\s+assessment",
        r"hazard\s+analysis", r"safety\s+audit",
        r"near\s+miss", r"lost\s+time",
        r"safety\s+record", r"injury\s+report",
    ],
    "geological_map": [
        r"geological\s+map", r"geologic\s+map",
        r"lithological\s+map", r"structural\s+map",
        r" alteration\s+map", r"geochemical\s+map",
        r"geophysical\s+map", r"topographic\s+map",
    ],
    "data_sheet": [
        r"data\s+sheet", r"spec\s+sheet", r"specification",
        r"datasheet", r"technical\s+specification",
        r"product\s+data", r"material\s+safety",
    ],
    "regulatory": [
        r"regulation", r"regulatory", r"compliance",
        r"permit\s+application", r"license\s+application",
        r"statute", r"act\s+of\s+", r"code\s+of\s+",
        r"ordinance", r"rule\s+and\s+regulation",
    ],
    "legal": [
        r"contract", r"agreement", r"terms\s+and\s+conditions",
        r"memorandum\s+of\s+understanding", r"MOU",
        r"joint\s+venture", r"earn[\-\s]in", r"royalty\s+agreement",
        r"mining\s+agreement", r"lease\s+agreement",
    ],
}

SECTION_TYPE_PATTERNS: Dict[str, List[str]] = {
    "introduction": [r"introduction", r"overview", r"background", r"scope", r"purpose"],
    "methodology": [r"methodology", r"methods", r"approach", r"technique", r"procedure", r"data\s+collection"],
    "results": [r"results", r"findings", r"analysis", r"discussion", r"interpretation", r"observations"],
    "conclusion": [r"conclusion", r"conclusions", r"summary", r"recommendation", r"recommendations", r"final\s+remarks"],
    "reference": [r"references", r"bibliography", r"literature\s+cited", r"works\s+cited", r"sources"],
}


@dataclass
class DocumentSection:
    heading: str
    content: str
    level: int
    page_number: int
    section_type: str = "other"


@dataclass
class TableData:
    headers: List[str]
    rows: List[List[str]]
    page_number: int = 0
    caption: str = ""


@dataclass
class FigureDescription:
    figure_number: str
    caption: str
    page_number: int = 0
    description: str = ""
    figure_type: str = "other"


@dataclass
class DocumentContent:
    doc_id: str = ""
    filename: str = ""
    file_type: str = ""
    content_text: str = ""
    page_count: int = 1
    word_count: int = 0
    char_count: int = 0
    sections: List[DocumentSection] = field(default_factory=list)
    key_findings: List[str] = field(default_factory=list)
    entities: Dict[str, List[str]] = field(default_factory=dict)
    tables: List[TableData] = field(default_factory=list)
    figures: List[FigureDescription] = field(default_factory=list)
    summary: str = ""
    key_terms: List[str] = field(default_factory=list)
    sentiment: str = "neutral"
    mining_relevance: float = 0.0
    related_topics: List[str] = field(default_factory=list)


class DocumentReader:
    """Intelligent document reading and understanding engine for mining AI."""

    def __init__(self, mining_vocabulary: Optional[Dict[str, List[str]]] = None):
        self.mining_vocabulary = mining_vocabulary or MINING_VOCABULARY
        self.doc_type_patterns = {
            dtype: [re.compile(p, re.IGNORECASE) for p in patterns]
            for dtype, patterns in DOC_TYPE_PATTERNS.items()
        }
        self.section_patterns = {
            stype: [re.compile(p, re.IGNORECASE) for p in patterns]
            for stype, patterns in SECTION_TYPE_PATTERNS.items()
        }
        self._all_mineral_set: Set[str] = set(
            m.lower() for m in self.mining_vocabulary.get("minerals", [])
        )
        self._all_equipment_set: Set[str] = set(
            e.lower() for e in self.mining_vocabulary.get("equipment", [])
        )
        self._all_process_set: Set[str] = set(
            p.lower() for p in self.mining_vocabulary.get("processes", [])
        )
        self._all_chemical_set: Set[str] = set(
            c.lower() for c in self.mining_vocabulary.get("chemicals", [])
        )
        self._all_regulation_set: Set[str] = set(
            r.lower() for r in self.mining_vocabulary.get("regulations", [])
        )
        self._all_geo_set: Set[str] = set(
            g.lower() for g in self.mining_vocabulary.get("geological_terms", [])
        )
        self._all_safety_set: Set[str] = set(
            s.lower() for s in self.mining_vocabulary.get("safety_terms", [])
        )
        self._build_term_indices()

    def _build_term_indices(self) -> None:
        self._mineral_terms_sorted = sorted(self._all_mineral_set, key=len, reverse=True)
        self._equipment_terms_sorted = sorted(self._all_equipment_set, key=len, reverse=True)
        self._process_terms_sorted = sorted(self._all_process_set, key=len, reverse=True)
        self._chemical_terms_sorted = sorted(self._all_chemical_set, key=len, reverse=True)
        self._regulation_terms_sorted = sorted(self._all_regulation_set, key=len, reverse=True)

    def _generate_doc_id(self, file_path: str) -> str:
        raw = f"{file_path}:{datetime.utcnow().isoformat()}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]

    def _count_words(self, text: str) -> int:
        return len(text.split())

    def _extract_terms_by_set(
        self, text: str, term_set: Set[str], terms_sorted: List[str]
    ) -> List[str]:
        text_lower = text.lower()
        found: List[str] = []
        seen: Set[str] = set()
        for term in terms_sorted:
            if term in seen:
                continue
            pattern = re.compile(r'\b' + re.escape(term) + r'\b', re.IGNORECASE)
            if pattern.search(text_lower):
                canonical = term
                found.append(canonical)
                seen.add(term)
        return found

    def _extract_terms_by_list(
        self, text: str, vocabulary_list: List[str]
    ) -> List[str]:
        term_set = set(t.lower() for t in vocabulary_list)
        sorted_terms = sorted(term_set, key=len, reverse=True)
        return self._extract_terms_by_set(text, term_set, sorted_terms)

    def extract_minerals(self, text: str) -> List[str]:
        return self._extract_terms_by_set(
            text, self._all_mineral_set, self._mineral_terms_sorted
        )

    def extract_equipment(self, text: str) -> List[str]:
        return self._extract_terms_by_set(
            text, self._all_equipment_set, self._equipment_terms_sorted
        )

    def extract_processes(self, text: str) -> List[str]:
        return self._extract_terms_by_set(
            text, self._all_process_set, self._process_terms_sorted
        )

    def extract_chemicals(self, text: str) -> List[str]:
        return self._extract_terms_by_set(
            text, self._all_chemical_set, self._chemical_terms_sorted
        )

    def extract_regulations(self, text: str) -> List[str]:
        return self._extract_terms_by_set(
            text, self._all_regulation_set, self._regulation_terms_sorted
        )

    def extract_locations(self, text: str) -> List[str]:
        locations: List[str] = []
        coord_pattern = re.compile(
            r'[-+]?\d{1,3}(?:\.\d+)?\s*[°]\s*[NSns]?\s*,?\s*[-+]?\d{1,3}(?:\.\d+)?\s*[°]\s*[EWew]?'
        )
        for match in coord_pattern.finditer(text):
            locations.append(match.group().strip())

        place_pattern = re.compile(
            r'\b(?:Mine|Mining|District|Region|Province|State|County|Shire|Basin|Range|Mount|Lake|River|Creek|Reef|Valley|Mountains?)\s+[A-Z][a-zA-Z\s]{1,30}'
        )
        for match in place_pattern.finditer(text):
            loc = match.group().strip()
            if loc not in locations:
                locations.append(loc)

        capitalized_pattern = re.compile(
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b'
        )
        common_words = {
            "the", "this", "that", "with", "from", "into", "were", "been",
            "being", "have", "has", "had", "will", "would", "could", "should",
            "may", "might", "shall", "can", "must", "not", "but", "and",
            "for", "are", "was", "were", "its", "our", "their", "your",
            "also", "such", "each", "than", "then", "when", "what", "how",
            "which", "where", "there", "here", "more", "most", "some",
            "only", "very", "after", "before", "during", "between",
            "table", "figure", "section", "chapter", "appendix", "reference",
            "report", "study", "analysis", "company", "project", "site",
            "resource", "reserve", "grade", "tonne", "ore", "rock",
        }
        for match in capitalized_pattern.finditer(text):
            place = match.group(1)
            if place.lower() not in common_words and place not in locations:
                locations.append(place)
            if len(locations) > 200:
                break
        return locations

    def extract_entities(self, text: str) -> Dict[str, List[str]]:
        return {
            "locations": self.extract_locations(text),
            "minerals": self.extract_minerals(text),
            "equipment": self.extract_equipment(text),
            "dates": self._extract_dates(text),
            "people": self._extract_people(text),
            "companies": self._extract_companies(text),
        }

    def _extract_dates(self, text: str) -> List[str]:
        dates: List[str] = []
        patterns = [
            re.compile(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'),
            re.compile(r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'),
            re.compile(
                r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)'
                r'\s+\d{1,2},?\s+\d{4}\b',
                re.IGNORECASE,
            ),
            re.compile(r'\b(?:Q[1-4]|[Ff]iscal\s+Year|FY)\s*\d{2,4}\b'),
            re.compile(r'\b(?:H[1-2]|first\s+half|second\s+half)\s+(?:of\s+)?\d{4}\b', re.IGNORECASE),
        ]
        for pat in patterns:
            for match in pat.finditer(text):
                d = match.group().strip()
                if d not in dates:
                    dates.append(d)
        return dates

    def _extract_people(self, text: str) -> List[str]:
        people: List[str] = []
        name_pattern = re.compile(
            r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.|Sir)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b'
        )
        for match in name_pattern.finditer(text):
            name = match.group().strip()
            if name not in people:
                people.append(name)

        title_pattern = re.compile(
            r'\b([A-Z][a-z]+\s+[A-Z][a-z]+)\s*,?\s*(?:CEO|CTO|CFO|COO|VP|Director|Manager|Engineer|Geologist|Analyst|Supervisor|Superintendent)',
        )
        for match in title_pattern.finditer(text):
            name = match.group(1).strip()
            if name not in people:
                people.append(name)
        return people[:100]

    def _extract_companies(self, text: str) -> List[str]:
        companies: List[str] = []
        suffixes = [
            "Inc", "Ltd", "LLC", "Corp", "Corporation", "Company",
            "Group", "Holdings", "Plc", "SA", "AG", "GmbH", "Pty",
            "Limited", "Enterprises", "Resources", "Mining", "Minerals",
            "Metals", "Energy", "Exploration", "Developments",
        ]
        for suffix in suffixes:
            pat = re.compile(
                r'\b([A-Z][A-Za-z\s&]{2,40})\s+' + re.escape(suffix) + r'\b'
            )
            for match in pat.finditer(text):
                company = match.group().strip()
                if len(company) > 3 and company not in companies:
                    companies.append(company)
        return companies

    def identify_document_type(self, text: str) -> str:
        scores: Dict[str, int] = {}
        text_lower = text.lower()
        for dtype, patterns in self.doc_type_patterns.items():
            count = 0
            for pat in patterns:
                count += len(pat.findall(text_lower))
            scores[dtype] = count
        if not scores or max(scores.values()) == 0:
            return "other"
        return max(scores, key=scores.get)

    def extract_key_findings(self, text: str, max_findings: int = 10) -> List[str]:
        findings: List[str] = []
        finding_patterns = [
            re.compile(
                r'(?:key\s+finding|finding|conclusion|result|discovery|recommendation)'
                r'[:\s]+(.+?)(?:\.|$)',
                re.IGNORECASE,
            ),
            re.compile(
                r'(?:study\s+(?:found|shows|indicates|reveals|demonstrates|confirms))'
                r'[:\s]+(.+?)(?:\.|$)',
                re.IGNORECASE,
            ),
            re.compile(
                r'(?:analysis\s+(?:shows|indicates|reveals|suggests|confirms))'
                r'[:\s]+(.+?)(?:\.|$)',
                re.IGNORECASE,
            ),
            re.compile(
                r'(?:it\s+was\s+(?:found|determined|observed|noted|discovered))'
                r'[:\s]+(.+?)(?:\.|$)',
                re.IGNORECASE,
            ),
            re.compile(
                r'(?:the\s+(?:results?|analysis|study|report|investigation))'
                r'\s+(?:shows?|indicates?|reveals?|suggests?|confirms?|demonstrates?)'
                r'[:\s]+(.+?)(?:\.|$)',
                re.IGNORECASE,
            ),
        ]
        for pat in finding_patterns:
            for match in pat.finditer(text):
                finding = match.group(1).strip()
                finding = re.sub(r'\s+', ' ', finding)
                if len(finding) > 20 and finding not in findings:
                    findings.append(finding)
                if len(findings) >= max_findings:
                    return findings

        sentences = re.split(r'(?<=[.!?])\s+', text)
        scoring_keywords = [
            "significant", "important", "notable", "key", "major",
            "critical", "essential", "demonstrate", "indicate", "reveal",
            "confirm", "suggest", "show", "result", "finding", "conclusion",
            "discovery", "breakthrough", "exceptional", "substantial",
        ]
        scored: List[Tuple[float, str]] = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 30 or len(sentence) > 500:
                continue
            score = 0.0
            s_lower = sentence.lower()
            for kw in scoring_keywords:
                if kw in s_lower:
                    score += 1.0
            has_number = bool(re.search(r'\d+\.?\d*\s*(?:%|tonnes?|MT|oz|g/t|%)', s_lower))
            if has_number:
                score += 2.0
            if any(f[0].lower() in s_lower[:50] for f in findings):
                continue
            if score > 0:
                scored.append((score, sentence))
        scored.sort(key=lambda x: x[0], reverse=True)
        for _, sent in scored[:max_findings]:
            if sent not in findings:
                findings.append(sent)
        return findings[:max_findings]

    def extract_tables_from_text(self, text: str) -> List[TableData]:
        tables: List[TableData] = []
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            pipe_count = line.count('|')
            tab_count = line.count('\t')
            if pipe_count >= 2 or tab_count >= 2:
                table_lines: List[str] = []
                j = i
                while j < len(lines):
                    curr = lines[j].strip()
                    if not curr:
                        if table_lines and len(table_lines) >= 2:
                            break
                        j += 1
                        continue
                    curr_pipes = curr.count('|')
                    curr_tabs = curr.count('\t')
                    if curr_pipes >= 2 or curr_tabs >= 2 or (table_lines and _looks_like_table_row(curr)):
                        table_lines.append(curr)
                        j += 1
                    else:
                        break
                if len(table_lines) >= 2:
                    table = self._parse_table_lines(table_lines)
                    if table and len(table.headers) >= 2:
                        tables.append(table)
                i = j
            else:
                i += 1
        return tables

    def _parse_table_lines(self, lines: List[str]) -> Optional[TableData]:
        if not lines:
            return None
        separator_pattern = re.compile(r'^[\s|:\-=\s]+$')
        headers: List[str] = []
        data_lines: List[str] = []
        for idx, line in enumerate(lines):
            if separator_pattern.match(line):
                continue
            cells = [c.strip() for c in re.split(r'[|\t]', line) if c.strip()]
            if not headers and all(c.isalpha() or c.replace(' ', '').isalpha() for c in cells if c):
                headers = cells
            elif cells:
                data_lines.append(line)
        if not headers and data_lines:
            first_cells = [c.strip() for c in re.split(r'[|\t]', data_lines[0]) if c.strip()]
            headers = first_cells
            data_lines = data_lines[1:]
        rows: List[List[str]] = []
        for line in data_lines:
            cells = [c.strip() for c in re.split(r'[|\t]', line) if c.strip()]
            if cells:
                while len(cells) < len(headers):
                    cells.append("")
                rows.append(cells[:len(headers)])
        return TableData(headers=headers, rows=rows, page_number=0)

    def _assess_sentiment(self, text: str) -> str:
        text_lower = text.lower()
        positive_words = [
            "successful", "significant", "excellent", "positive", "increase",
            "improved", "optimal", "efficient", "favorable", "profitable",
            "high grade", "substantial", "promising", "strong", "growth",
            "exceeds", "outperform", "robust", "sustainable", "innovative",
            "breakthrough", "exceptional", "remarkable", "abundant", "rich",
        ]
        negative_words = [
            "failure", "decline", "decrease", "loss", "negative", "risk",
            "hazard", "contamination", "degradation", "damage", "accident",
            "incident", "pollution", "deficit", "poor", "insufficient",
            "concern", "threat", "violation", "noncompliance", "breach",
            "impairment", "deterioration", "shortfall", "deficiency",
        ]
        pos_count = sum(1 for w in positive_words if w in text_lower)
        neg_count = sum(1 for w in negative_words if w in text_lower)
        if pos_count > neg_count * 1.5:
            return "positive"
        elif neg_count > pos_count * 1.5:
            return "negative"
        return "neutral"

    def generate_executive_summary(self, text: str, max_length: int = 500) -> str:
        if len(text) <= max_length:
            return text.strip()
        sentences = re.split(r'(?<=[.!?])\s+', text)
        scoring_keywords = [
            "key", "main", "primary", "significant", "important",
            "conclusion", "result", "finding", "recommendation",
            "summary", "overall", "total", "aggregate", "project",
            "resource", "reserve", "grade", "tonnage", "recovery",
        ]
        scored: List[Tuple[float, int, str]] = []
        for idx, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if len(sentence) < 20:
                continue
            score = 0.0
            s_lower = sentence.lower()
            for kw in scoring_keywords:
                if kw in s_lower:
                    score += 1.0
            score += max(0, 5 - idx * 0.1)
            has_number = bool(re.search(r'\d+\.?\d*', sentence))
            if has_number:
                score += 1.5
            scored.append((score, idx, sentence))
        scored.sort(key=lambda x: x[0], reverse=True)
        scored.sort(key=lambda x: x[1])
        summary_parts: List[str] = []
        current_length = 0
        for _, _, sentence in scored:
            if current_length + len(sentence) + 1 <= max_length:
                summary_parts.append(sentence)
                current_length += len(sentence) + 1
            if current_length >= max_length:
                break
        if not summary_parts:
            truncated = text[:max_length]
            last_period = truncated.rfind('.')
            if last_period > max_length * 0.5:
                return truncated[:last_period + 1].strip()
            return truncated.strip() + "..."
        return " ".join(summary_parts)

    def assess_mining_relevance(self, text: str) -> float:
        text_lower = text.lower()
        total_hits = 0
        all_terms = (
            list(self._all_mineral_set)
            + list(self._all_equipment_set)
            + list(self._all_process_set)
            + list(self._all_chemical_set)
            + list(self._all_geo_set)
        )
        for term in all_terms:
            if re.search(r'\b' + re.escape(term) + r'\b', text_lower):
                total_hits += 1
        if total_hits == 0:
            return 0.0
        score = min(1.0, total_hits / 30.0)
        doc_type = self.identify_document_type(text)
        if doc_type in ("feasibility_study", "technical_report", "geological_map"):
            score = min(1.0, score + 0.2)
        elif doc_type in ("environmental_impact", "safety_report"):
            score = min(1.0, score + 0.1)
        words = text_lower.split()
        if not words:
            return score
        mining_ratio = sum(
            1 for w in words
            if any(
                re.search(r'\b' + re.escape(term) + r'\b', w)
                for term in all_terms[:200]
            )
        ) / len(words)
        score = min(1.0, score + mining_ratio * 5)
        return round(score, 3)

    def find_related_topics(self, text: str) -> List[str]:
        topics: List[str] = []
        topic_checks: List[Tuple[str, List[str]]] = [
            ("mineral exploration", ["exploration", "drill", "assay", "resource", "geological survey"]),
            ("mining operations", ["mining", "excavation", "haul", "pit", "underground"]),
            ("metallurgical processing", ["crush", "mill", "flotation", "leach", "smelt", "refin"]),
            ("environmental management", ["environment", "reclamation", "tailings", "water quality", "air quality"]),
            ("mine safety", ["safety", "hazard", "risk", "PPE", "incident", "accident"]),
            ("financial analysis", ["cost", "revenue", "NPV", "IRR", "capital", "operating"]),
            ("geotechnical engineering", ["geotechnical", "slope stability", "ground control", "rock mechanics"]),
            ("water management", ["water", "dewatering", "groundwater", "surface water", "drainage"]),
            ("community relations", ["community", "indigenous", "stakeholder", "consultation", "social"]),
            ("regulatory compliance", ["permit", "regulation", "compliance", "license", "EPA"]),
            ("mine closure", ["closure", "decommissioning", "reclamation", "post-closure"]),
            ("drilling and blasting", ["drill", "blast", "explosive", "detonation"]),
            ("material handling", ["conveyor", "haul", "transport", "stockpile"]),
            ("resource estimation", ["resource", "reserve", "block model", "kriging", "geostatistic"]),
        ]
        text_lower = text.lower()
        for topic_name, keywords in topic_checks:
            hits = sum(1 for kw in keywords if kw in text_lower)
            if hits >= 2:
                topics.append(topic_name)
        return topics

    def detect_anomalies_in_data(self, text: str) -> List[str]:
        anomalies: List[str] = []
        tables = self.extract_tables_from_text(text)
        for t_idx, table in enumerate(tables):
            for col_idx, header in enumerate(table.headers):
                numeric_values: List[Tuple[int, float]] = []
                for row_idx, row in enumerate(table.rows):
                    if col_idx < len(row):
                        val_str = row[col_idx].strip().replace(',', '').replace('%', '').replace('$', '')
                        try:
                            val = float(val_str)
                            numeric_values.append((row_idx, val))
                        except ValueError:
                            pass
                if len(numeric_values) < 3:
                    continue
                values = [v for _, v in numeric_values]
                mean_val = sum(values) / len(values)
                if mean_val == 0:
                    continue
                variance = sum((v - mean_val) ** 2 for v in values) / len(values)
                std_dev = math.sqrt(variance)
                for row_idx, val in numeric_values:
                    if std_dev > 0 and abs(val - mean_val) > 3 * std_dev:
                        anomalies.append(
                            f"Table {t_idx + 1}, column '{header}', row {row_idx + 1}: "
                            f"value {val} is >3 std devs from mean ({mean_val:.2f})"
                        )
                if len(values) >= 2:
                    for idx in range(1, len(values)):
                        if values[idx - 1] != 0:
                            ratio = values[idx] / values[idx - 1]
                            if ratio > 10 or ratio < 0.1:
                                anomalies.append(
                                    f"Table {t_idx + 1}, column '{header}': "
                                    f"large jump from {values[idx - 1]} to {values[idx]} "
                                    f"(ratio: {ratio:.2f})"
                                )
        numeric_pattern = re.compile(
            r'(\d[\d,]*\.?\d*)\s*(tonnes?|MT|oz|g/t|ppm|ppb|%)\b', re.IGNORECASE
        )
        all_numbers: List[Tuple[str, float]] = []
        for match in numeric_pattern.finditer(text):
            unit = match.group(2).lower()
            try:
                num_str = match.group(1).replace(',', '')
                val = float(num_str)
                all_numbers.append((unit, val))
            except ValueError:
                pass
        units_seen = set(u for u, _ in all_numbers)
        for unit in units_seen:
            unit_vals = [v for u, v in all_numbers if u == unit]
            if len(unit_vals) < 3:
                continue
            mean_val = sum(unit_vals) / len(unit_vals)
            if mean_val == 0:
                continue
            for val in unit_vals:
                if abs(val - mean_val) > 2 * (sum(abs(v - mean_val) for v in unit_vals) / len(unit_vals)):
                    anomalies.append(
                        f"Numeric value {val} {unit} appears anomalous vs document mean of {mean_val:.2f} {unit}"
                    )
        return anomalies

    def _classify_section_type(self, heading: str) -> str:
        heading_lower = heading.lower().strip()
        for stype, patterns in self.section_patterns.items():
            for pat in patterns:
                if pat.search(heading_lower):
                    return stype
        return "other"

    def _parse_sections_from_text(self, text: str, page_number: int = 1) -> List[DocumentSection]:
        sections: List[DocumentSection] = []
        heading_pattern = re.compile(
            r'^(#{1,6})\s+(.+)$|'
            r'^([A-Z][A-Za-z0-9\s\-:]{2,80})$',
            re.MULTILINE,
        )
        matches = list(heading_pattern.finditer(text))
        if not matches:
            return sections
        for idx, match in enumerate(matches):
            if match.group(1):
                level = len(match.group(1))
                heading = match.group(2).strip()
            else:
                heading = match.group(3).strip()
                if len(heading) > 100:
                    continue
                words = heading.split()
                if len(words) <= 2:
                    level = 1
                elif len(words) <= 4:
                    level = 2
                else:
                    level = 3
            start = match.end()
            end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            if not content:
                content = ""
            section_type = self._classify_section_type(heading)
            sections.append(DocumentSection(
                heading=heading,
                content=content,
                level=level,
                page_number=page_number,
                section_type=section_type,
            ))
        return sections

    def _extract_figures_from_text(self, text: str) -> List[FigureDescription]:
        figures: List[FigureDescription] = []
        figure_pattern = re.compile(
            r'(?:Figure|Fig\.?|Image|Diagram|Chart|Photo|Plate)\s*(\d+(?:\.\d+)?)'
            r'[:\.\s]+(.+?)(?:\.|$)',
            re.IGNORECASE | re.MULTILINE,
        )
        for match in figure_pattern.finditer(text):
            fig_num = match.group(1)
            caption = match.group(2).strip()
            fig_type = "other"
            caption_lower = caption.lower()
            if any(w in caption_lower for w in ["chart", "graph", "plot", "bar", "histogram", "pie"]):
                fig_type = "chart"
            elif any(w in caption_lower for w in ["map", "plan", "aerial", "satellite"]):
                fig_type = "map"
            elif any(w in caption_lower for w in ["photo", "photograph", "image", "picture"]):
                fig_type = "photo"
            elif any(w in caption_lower for w in ["diagram", "schematic", "flowchart", "flow"]):
                fig_type = "diagram"
            elif any(w in caption_lower for w in ["table", "tabular"]):
                fig_type = "table"
            figures.append(FigureDescription(
                figure_number=fig_num,
                caption=caption,
                page_number=0,
                description=caption,
                figure_type=fig_type,
            ))
        return figures

    def _build_content_object(
        self, file_path: str, text: str, page_count: int = 1, file_type: str = "txt"
    ) -> DocumentContent:
        doc_id = self._generate_doc_id(file_path)
        filename = os.path.basename(file_path)
        word_count = self._count_words(text)
        char_count = len(text)
        sections = self._parse_sections_from_text(text)
        key_findings = self.extract_key_findings(text)
        entities = self.extract_entities(text)
        tables = self.extract_tables_from_text(text)
        figures = self._extract_figures_from_text(text)
        doc_type = self.identify_document_type(text)
        summary = self.generate_executive_summary(text)
        mining_relevance = self.assess_mining_relevance(text)
        related_topics = self.find_related_topics(text)
        sentiment = self._assess_sentiment(text)
        key_terms = self._extract_key_terms(text)
        return DocumentContent(
            doc_id=doc_id,
            filename=filename,
            file_type=file_type,
            content_text=text,
            page_count=page_count,
            word_count=word_count,
            char_count=char_count,
            sections=sections,
            key_findings=key_findings,
            entities=entities,
            tables=tables,
            figures=figures,
            summary=summary,
            key_terms=key_terms,
            sentiment=sentiment,
            mining_relevance=mining_relevance,
            related_topics=related_topics,
        )

    def _extract_key_terms(self, text: str) -> List[str]:
        text_lower = text.lower()
        words = re.findall(r'\b[a-z]{3,}\b', text_lower)
        word_freq: Dict[str, int] = {}
        for w in words:
            word_freq[w] = word_freq.get(w, 0) + 1
        mining_terms = set()
        for category_terms in self.mining_vocabulary.values():
            for term in category_terms:
                mining_terms.add(term.lower())
        stop_words = {
            "the", "and", "for", "are", "but", "not", "you", "all", "can",
            "had", "her", "was", "one", "our", "out", "has", "his", "how",
            "its", "may", "new", "now", "old", "see", "way", "who", "did",
            "get", "let", "say", "she", "too", "use", "this", "that", "with",
            "have", "from", "they", "been", "said", "each", "which",
            "their", "time", "will", "way", "about", "many", "then",
            "them", "would", "write", "like", "so", "these", "her",
            "long", "make", "thing", "see", "him", "two", "has", "look",
            "more", "day", "could", "go", "come", "did", "my", "number",
            "sound", "no", "most", "people", "my", "over", "know",
            "water", "than", "call", "first", "who", "may", "down",
            "side", "been", "now", "find", "head", "stand", "own",
            "page", "should", "country", "found", "answer", "school",
            "grow", "still", "learn", "should", "America", "world",
        }
        key_terms: List[Tuple[str, float]] = []
        for word, freq in word_freq.items():
            if word in stop_words or len(word) < 4:
                continue
            is_mining = word in mining_terms
            if is_mining:
                score = freq * 3.0
            elif freq >= 3:
                score = freq * 1.0
            else:
                continue
            key_terms.append((word, score))
        key_terms.sort(key=lambda x: x[1], reverse=True)
        return [term for term, _ in key_terms[:30]]

    def read_pdf(self, file_path: str) -> DocumentContent:
        try:
            return self._read_pdf_pypdf(file_path)
        except ImportError:
            logger.info("pypdf not available, using basic binary parsing for PDF")
            return self._read_pdf_basic(file_path)
        except Exception as e:
            logger.warning(f"pypdf failed ({e}), falling back to basic parsing")
            return self._read_pdf_basic(file_path)

    def _read_pdf_pypdf(self, file_path: str) -> DocumentContent:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        page_count = len(reader.pages)
        text_parts: List[str] = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        full_text = "\n\n".join(text_parts)
        return self._build_content_object(file_path, full_text, page_count, "pdf")

    def _read_pdf_basic(self, file_path: str) -> DocumentContent:
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            text = self._extract_text_from_pdf_binary(content)
            page_count = content.count(b'/Type /Page') or content.count(b'/Type/Page') or 1
            if isinstance(page_count, int) and page_count < 1:
                page_count = 1
            return self._build_content_object(file_path, text, page_count, "pdf")
        except Exception as e:
            logger.error(f"Failed to read PDF {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading PDF: {e}]", 1, "pdf")

    def _extract_text_from_pdf_binary(self, pdf_bytes: bytes) -> str:
        text_parts: List[str] = []
        text_pattern = re.compile(rb'\(([^)]+)\)\s*Tj')
        text_array_pattern = re.compile(rb'\[(.*?)\]\s*TJ', re.DOTALL)
        for match in text_pattern.finditer(pdf_bytes):
            try:
                decoded = match.group(1).decode('latin-1', errors='replace')
                text_parts.append(decoded)
            except Exception:
                pass
        for match in text_array_pattern.finditer(pdf_bytes):
            try:
                parts = re.findall(rb'\(([^)]+)\)', match.group(1))
                line = ''.join(
                    p.decode('latin-1', errors='replace') for p in parts
                )
                text_parts.append(line)
            except Exception:
                pass
        return '\n'.join(text_parts)

    def read_docx(self, file_path: str) -> DocumentContent:
        try:
            return self._read_docx_python_docx(file_path)
        except ImportError:
            logger.info("python-docx not available, using XML parsing for DOCX")
            return self._read_docx_xml(file_path)
        except Exception as e:
            logger.warning(f"python-docx failed ({e}), falling back to XML parsing")
            return self._read_docx_xml(file_path)

    def _read_docx_python_docx(self, file_path: str) -> DocumentContent:
        from docx import Document
        doc = Document(file_path)
        text_parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style and para.style.name.startswith('Heading'):
                    level_match = re.search(r'\d', para.style.name)
                    level = int(level_match.group()) if level_match else 1
                    text_parts.append('#' * level + ' ' + para.text)
                else:
                    text_parts.append(para.text)
        for table in doc.tables:
            rows_text: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(' | '.join(cells))
            if rows_text:
                text_parts.append('\n'.join(rows_text))
        full_text = '\n\n'.join(text_parts)
        return self._build_content_object(file_path, full_text, 1, "docx")

    def _read_docx_xml(self, file_path: str) -> DocumentContent:
        import zipfile
        text_parts: List[str] = []
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                if 'word/document.xml' in z.namelist():
                    with z.open('word/document.xml') as f:
                        xml_content = f.read().decode('utf-8', errors='replace')
                    text_pattern = re.compile(r'<w:t[^>]*>(.*?)</w:t>', re.DOTALL)
                    for match in text_pattern.finditer(xml_content):
                        text = match.group(1)
                        text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
                        text_parts.append(text)
                    paragraph_pattern = re.compile(r'<w:p[^>]*>(.*?)</w:p>', re.DOTALL)
                    paragraphs: List[str] = []
                    for match in paragraph_pattern.finditer(xml_content):
                        para_xml = match.group(1)
                        para_texts = text_pattern.findall(para_xml)
                        para_text = ''.join(para_texts).strip()
                        if para_text:
                            heading_match = re.search(r'<w:pStyle[^>]*w:val="Heading(\d)"', para_xml)
                            if heading_match:
                                level = int(heading_match.group(1))
                                para_text = '#' * level + ' ' + para_text
                            paragraphs.append(para_text)
                    text_parts = paragraphs
        except Exception as e:
            logger.error(f"Failed to parse DOCX XML {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading DOCX: {e}]", 1, "docx")
        full_text = '\n\n'.join(text_parts)
        return self._build_content_object(file_path, full_text, 1, "docx")

    def read_xlsx(self, file_path: str) -> DocumentContent:
        try:
            return self._read_xlsx_openpyxl(file_path)
        except ImportError:
            logger.info("openpyxl not available, using ZIP/XML parsing for XLSX")
            return self._read_xlsx_zip(file_path)
        except Exception as e:
            logger.warning(f"openpyxl failed ({e}), falling back to ZIP parsing")
            return self._read_xlsx_zip(file_path)

    def _read_xlsx_openpyxl(self, file_path: str) -> DocumentContent:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        text_parts: List[str] = []
        tables: List[TableData] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"## Sheet: {sheet_name}")
            all_rows: List[List[str]] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(c.strip() for c in cells):
                    all_rows.append(cells)
            if all_rows:
                headers = all_rows[0]
                data_rows = all_rows[1:]
                text_parts.append(" | ".join(headers))
                for row in data_rows:
                    text_parts.append(" | ".join(row))
                tables.append(TableData(
                    headers=headers,
                    rows=data_rows,
                    page_number=0,
                    caption=f"Sheet: {sheet_name}",
                ))
        wb.close()
        full_text = '\n\n'.join(text_parts)
        doc = self._build_content_object(file_path, full_text, 1, "xlsx")
        doc.tables = tables
        return doc

    def _read_xlsx_zip(self, file_path: str) -> DocumentContent:
        import zipfile
        text_parts: List[str] = []
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                sheet_files = [f for f in z.namelist() if f.startswith('xl/worksheets/sheet')]
                shared_strings: List[str] = []
                if 'xl/sharedStrings.xml' in z.namelist():
                    with z.open('xl/sharedStrings.xml') as f:
                        ss_xml = f.read().decode('utf-8', errors='replace')
                    ss_pattern = re.compile(r'<t[^>]*>(.*?)</t>', re.DOTALL)
                    shared_strings = ss_pattern.findall(ss_xml)
                    shared_strings = [
                        s.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
                        for s in shared_strings
                    ]
                for sheet_file in sorted(sheet_files):
                    with z.open(sheet_file) as f:
                        sheet_xml = f.read().decode('utf-8', errors='replace')
                    row_pattern = re.compile(r'<row[^>]*>(.*?)</row>', re.DOTALL)
                    cell_pattern = re.compile(
                        r'<c[^>]*(?:t="s")?[^>]*>(?:<v>(.*?)</v>)?</c>', re.DOTALL
                    )
                    rows_text: List[str] = []
                    for row_match in row_pattern.finditer(sheet_xml):
                        row_xml = row_match.group(1)
                        cells = cell_pattern.findall(row_xml)
                        row_values: List[str] = []
                        for cell_val in cells:
                            if cell_val is None:
                                row_values.append("")
                            else:
                                try:
                                    idx = int(cell_val)
                                    if idx < len(shared_strings):
                                        row_values.append(shared_strings[idx])
                                    else:
                                        row_values.append(cell_val)
                                except ValueError:
                                    row_values.append(cell_val)
                        if any(v.strip() for v in row_values):
                            rows_text.append(" | ".join(row_values))
                    sheet_name = os.path.basename(sheet_file).replace('.xml', '')
                    text_parts.append(f"## Sheet: {sheet_name}")
                    text_parts.extend(rows_text)
        except Exception as e:
            logger.error(f"Failed to parse XLSX zip {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading XLSX: {e}]", 1, "xlsx")
        full_text = '\n\n'.join(text_parts)
        return self._build_content_object(file_path, full_text, 1, "xlsx")

    def read_csv(self, file_path: str) -> DocumentContent:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
        except Exception as e:
            logger.error(f"Failed to read CSV {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading CSV: {e}]", 1, "csv")
        lines = content.strip().split('\n')
        if not lines:
            return self._build_content_object(file_path, "", 1, "csv")
        delimiter = self._detect_csv_delimiter(lines[0])
        headers = [h.strip().strip('"') for h in lines[0].split(delimiter)]
        rows: List[List[str]] = []
        for line in lines[1:]:
            if line.strip():
                cells = [c.strip().strip('"') for c in line.split(delimiter)]
                rows.append(cells)
        table_text = " | ".join(headers) + "\n"
        for row in rows:
            table_text += " | ".join(row) + "\n"
        doc = self._build_content_object(file_path, content, 1, "csv")
        doc.tables = [TableData(headers=headers, rows=rows, page_number=0, caption=os.path.basename(file_path))]
        return doc

    def _detect_csv_delimiter(self, header_line: str) -> str:
        delimiters = [',', '\t', ';', '|']
        best_delim = ','
        best_count = 0
        for delim in delimiters:
            count = header_line.count(delim)
            if count > best_count:
                best_count = count
                best_delim = delim
        return best_delim

    def read_txt(self, file_path: str) -> DocumentContent:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
        except Exception as e:
            logger.error(f"Failed to read TXT {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading TXT: {e}]", 1, "txt")
        page_count = max(1, len(content) // 3000)
        return self._build_content_object(file_path, content, page_count, "txt")

    def read_json(self, file_path: str) -> DocumentContent:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                raw = f.read()
        except Exception as e:
            logger.error(f"Failed to read JSON {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading JSON: {e}]", 1, "json")
        try:
            data = json.loads(raw)
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {file_path}: {e}")
            return self._build_content_object(file_path, raw, 1, "json")
        readable_text = self._flatten_json_to_text(data, max_depth=8)
        doc = self._build_content_object(file_path, readable_text, 1, "json")
        return doc

    def _flatten_json_to_text(self, obj: Any, max_depth: int = 8, current_depth: int = 0) -> str:
        if current_depth > max_depth:
            return str(obj)
        if isinstance(obj, str):
            return obj
        elif isinstance(obj, (int, float, bool)):
            return str(obj)
        elif obj is None:
            return ""
        elif isinstance(obj, list):
            parts: List[str] = []
            for item in obj:
                parts.append(self._flatten_json_to_text(item, max_depth, current_depth + 1))
            return '\n'.join(parts)
        elif isinstance(obj, dict):
            parts = []
            for key, value in obj.items():
                val_text = self._flatten_json_to_text(value, max_depth, current_depth + 1)
                if val_text:
                    parts.append(f"{key}: {val_text}")
            return '\n'.join(parts)
        return str(obj)

    def read_image_info(self, file_path: str) -> DocumentContent:
        try:
            with open(file_path, 'rb') as f:
                header = f.read(32)
            image_info = self._parse_image_header(header, file_path)
            readable_text = (
                f"Image File: {os.path.basename(file_path)}\n"
                f"Format: {image_info['format']}\n"
                f"Width: {image_info.get('width', 'unknown')} pixels\n"
                f"Height: {image_info.get('height', 'unknown')} pixels\n"
                f"Color Depth: {image_info.get('color_depth', 'unknown')} bits\n"
                f"File Size: {image_info.get('file_size', 'unknown')} bytes\n"
            )
            if image_info.get('format') == 'PNG':
                readable_text += f"Color Type: {image_info.get('color_type', 'unknown')}\n"
                readable_text += f"Compression: {image_info.get('compression', 'unknown')}\n"
            elif image_info.get('format') == 'JPEG':
                readable_text += f"Color Space: {image_info.get('color_space', 'unknown')}\n"
                readable_text += f"Bits Per Component: {image_info.get('bits_per_component', 'unknown')}\n"
            doc = self._build_content_object(file_path, readable_text, 1, image_info['format'].lower())
            doc.figures = [FigureDescription(
                figure_number="1",
                caption=f"Image: {os.path.basename(file_path)}",
                page_number=1,
                description=readable_text,
                figure_type="photo",
            )]
            return doc
        except Exception as e:
            logger.error(f"Failed to read image info {file_path}: {e}")
            return self._build_content_object(file_path, f"[Error reading image: {e}]", 1, "image")

    def _parse_image_header(self, header: bytes, file_path: str) -> Dict[str, Any]:
        info: Dict[str, Any] = {
            "format": "unknown",
            "width": None,
            "height": None,
            "color_depth": None,
            "file_size": os.path.getsize(file_path),
        }
        if header[:8] == b'\x89PNG\r\n\x1a\n':
            info["format"] = "PNG"
            if len(header) >= 24:
                try:
                    width = int.from_bytes(header[16:20], 'big')
                    height = int.from_bytes(header[20:24], 'big')
                    info["width"] = width
                    info["height"] = height
                except Exception:
                    pass
            if len(header) >= 26:
                bit_depth = header[24]
                color_type = header[25]
                info["color_depth"] = bit_depth * (3 if color_type == 2 else 4 if color_type == 6 else 2 if color_type == 4 else 1)
                ct_names = {0: "Grayscale", 2: "RGB", 3: "Indexed", 4: "Grayscale+Alpha", 6: "RGBA"}
                info["color_type"] = ct_names.get(color_type, f"Unknown ({color_type})")
                info["compression"] = "Deflate"
        elif header[:2] == b'\xff\xd8':
            info["format"] = "JPEG"
            offset = 2
            while offset < len(header) - 1:
                if header[offset] != 0xFF:
                    break
                marker = header[offset + 1]
                if marker in (0xC0, 0xC1, 0xC2):
                    if offset + 9 < len(header):
                        try:
                            info["color_depth"] = header[offset + 4]
                            info["height"] = int.from_bytes(header[offset + 5:offset + 7], 'big')
                            info["width"] = int.from_bytes(header[offset + 7:offset + 9], 'big')
                            info["color_space"] = "RGB" if header[offset + 9] == 3 else "Grayscale" if header[offset + 9] == 1 else f"Type {header[offset + 9]}"
                            info["bits_per_component"] = header[offset + 4]
                        except Exception:
                            pass
                    break
                offset += 2
                if marker == 0xD9:
                    break
                if offset + 1 < len(header):
                    try:
                        seg_len = int.from_bytes(header[offset:offset + 2], 'big')
                        offset += seg_len
                    except Exception:
                        break
                else:
                    break
        elif header[:6] in (b'GIF87a', b'GIF89a'):
            info["format"] = "GIF"
            if len(header) >= 10:
                try:
                    info["width"] = int.from_bytes(header[6:8], 'little')
                    info["height"] = int.from_bytes(header[8:10], 'little')
                    info["color_depth"] = (header[10] & 0x07) + 1
                except Exception:
                    pass
        elif header[:4] == b'BM':
            info["format"] = "BMP"
            if len(header) >= 26:
                try:
                    info["width"] = int.from_bytes(header[18:22], 'little')
                    info["height"] = abs(int.from_bytes(header[22:26], 'little'))
                    if len(header) >= 28:
                        info["color_depth"] = int.from_bytes(header[28:30], 'little')
                except Exception:
                    pass
        elif header[:4] == b'\x00\x00\x01\x00':
            info["format"] = "ICO"
            if len(header) >= 6:
                try:
                    info["width"] = header[4] or 256
                    info["height"] = header[5] or 256
                except Exception:
                    pass
        elif header[:4] == b'RIFF' and header[8:12] == b'WEBP':
            info["format"] = "WEBP"
        else:
            ext = os.path.splitext(file_path)[1].lower()
            format_map = {
                ".jpg": "JPEG", ".jpeg": "JPEG", ".png": "PNG",
                ".gif": "GIF", ".bmp": "BMP", ".tiff": "TIFF",
                ".tif": "TIFF", ".webp": "WEBP", ".svg": "SVG",
                ".ico": "ICO", ".raw": "RAW", ".cr2": "CR2",
                ".nef": "NEF", ".arw": "ARW", ".dng": "DNG",
            }
            info["format"] = format_map.get(ext, "Unknown")
        return info

    def read_document(self, file_path: str) -> DocumentContent:
        if not os.path.exists(file_path):
            return DocumentContent(
                doc_id=self._generate_doc_id(file_path),
                filename=os.path.basename(file_path),
                file_type="unknown",
                content_text=f"[Error: File not found: {file_path}]",
            )
        ext = os.path.splitext(file_path)[1].lower()
        reader_map = {
            ".pdf": self.read_pdf,
            ".docx": self.read_docx,
            ".doc": self.read_docx,
            ".xlsx": self.read_xlsx,
            ".xls": self.read_xlsx,
            ".csv": self.read_csv,
            ".txt": self.read_txt,
            ".text": self.read_txt,
            ".log": self.read_txt,
            ".json": self.read_json,
            ".png": self.read_image_info,
            ".jpg": self.read_image_info,
            ".jpeg": self.read_image_info,
            ".gif": self.read_image_info,
            ".bmp": self.read_image_info,
            ".tiff": self.read_image_info,
            ".tif": self.read_image_info,
            ".webp": self.read_image_info,
            ".svg": self.read_image_info,
            ".ico": self.read_image_info,
        }
        reader = reader_map.get(ext)
        if reader:
            return reader(file_path)
        try:
            return self.read_txt(file_path)
        except Exception as e:
            logger.error(f"Failed to read {file_path}: {e}")
            return DocumentContent(
                doc_id=self._generate_doc_id(file_path),
                filename=os.path.basename(file_path),
                file_type=ext.lstrip("."),
                content_text=f"[Error: Could not read file: {e}]",
            )

    def read_and_understand(self, file_path: str) -> DocumentContent:
        doc = self.read_document(file_path)
        if doc.content_text.startswith("[Error:"):
            return doc
        doc.mining_relevance = self.assess_mining_relevance(doc.content_text)
        doc.sentiment = self._assess_sentiment(doc.content_text)
        doc.related_topics = self.find_related_topics(doc.content_text)
        doc.key_terms = self._extract_key_terms(doc.content_text)
        doc.entities = self.extract_entities(doc.content_text)
        doc.key_findings = self.extract_key_findings(doc.content_text)
        doc.tables = self.extract_tables_from_text(doc.content_text)
        doc.figures = self._extract_figures_from_text(doc.content_text)
        doc.sections = self._parse_sections_from_text(doc.content_text)
        doc.summary = self.generate_executive_summary(doc.content_text)
        anomalies = self.detect_anomalies_in_data(doc.content_text)
        if anomalies:
            doc.key_findings.extend(
                [f"[ANOMALY] {a}" for a in anomalies[:5]]
            )
        return doc

    def find_references_to_document(
        self, doc_id: str, all_documents: Dict[str, DocumentContent]
    ) -> List[str]:
        referencing_docs: List[str] = []
        target_doc = all_documents.get(doc_id)
        if not target_doc:
            return referencing_docs
        target_terms = set(t.lower() for t in target_doc.key_terms[:15])
        target_entities = set()
        for entity_list in target_doc.entities.values():
            for entity in entity_list:
                target_entities.add(entity.lower())
        for other_id, other_doc in all_documents.items():
            if other_id == doc_id:
                continue
            other_text = other_doc.content_text.lower()
            matches = 0
            for term in target_terms:
                if re.search(r'\b' + re.escape(term) + r'\b', other_text):
                    matches += 1
            for entity in target_entities:
                if re.search(r'\b' + re.escape(entity) + r'\b', other_text):
                    matches += 1
            if matches >= 3:
                referencing_docs.append(other_id)
        return referencing_docs

    def build_reference_graph(
        self, documents: Dict[str, DocumentContent]
    ) -> Dict[str, List[str]]:
        graph: Dict[str, List[str]] = {}
        for doc_id in documents:
            graph[doc_id] = self.find_references_to_document(doc_id, documents)
        return graph

    def suggest_related_documents(
        self,
        doc_id: str,
        all_documents: Dict[str, DocumentContent],
        limit: int = 5,
    ) -> List[Tuple[str, float]]:
        target_doc = all_documents.get(doc_id)
        if not target_doc:
            return []
        target_words = set(target_doc.content_text.lower().split())
        target_terms = set(t.lower() for t in target_doc.key_terms)
        target_entities = set()
        for entity_list in target_doc.entities.values():
            for entity in entity_list:
                target_entities.add(entity.lower())
        target_topics = set(t.lower() for t in target_doc.related_topics)
        scored: List[Tuple[str, float]] = []
        for other_id, other_doc in all_documents.items():
            if other_id == doc_id:
                continue
            score = 0.0
            other_words = set(other_doc.content_text.lower().split())
            if target_words and other_words:
                overlap = len(target_words & other_words)
                total = len(target_words | other_words)
                if total > 0:
                    score += (overlap / total) * 0.3
            other_terms = set(t.lower() for t in other_doc.key_terms)
            term_overlap = len(target_terms & other_terms)
            score += min(1.0, term_overlap / 10.0) * 0.25
            other_entities = set()
            for elist in other_doc.entities.values():
                for e in elist:
                    other_entities.add(e.lower())
            entity_overlap = len(target_entities & other_entities)
            score += min(1.0, entity_overlap / 10.0) * 0.25
            other_topics = set(t.lower() for t in other_doc.related_topics)
            topic_overlap = len(target_topics & other_topics)
            if target_topics:
                score += (topic_overlap / len(target_topics)) * 0.2
            if target_doc.file_type == other_doc.file_type:
                score += 0.05
            if target_doc.sentiment == other_doc.sentiment:
                score += 0.05
            if target_doc.doc_type == getattr(other_doc, 'doc_type', ''):
                score += 0.1
            scored.append((other_id, round(score, 4)))
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored[:limit]

    @property
    def doc_type(self) -> str:
        return "DocumentReader"
