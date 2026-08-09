import os
import io
import logging
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

logger = logging.getLogger("ai_os.document_service")

class DocumentService:
    """
    Document Service handles generation, extraction, and conversion of multiple document formats:
    TXT, PDF, DOCX, XLSX, CSV.
    """
    def __init__(self, minio_client: Optional[Any] = None):
        self.minio_client = minio_client

    def create_txt_from_content(self, filename: str, content: str) -> bytes:
        """
        Creates a raw TXT byte stream from string content.
        """
        return content.encode("utf-8")

    def create_pdf_from_content(self, title: str, content: str, table_data: Optional[List[List[Any]]] = None) -> bytes:
        """
        Generates a styled PDF report using ReportLab.
        Supports paragraphs, spacers, titles, and data tables.
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=40, leftMargin=40,
                                topMargin=40, bottomMargin=40)
        
        styles = getSampleStyleSheet()
        
        # Premium Custom Styles
        title_style = ParagraphStyle(
            name='PremiumTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=24,
            leading=28,
            textColor=colors.HexColor('#0F172A'), # Charcoal / Navy slate
            spaceAfter=15
        )
        
        body_style = ParagraphStyle(
            name='PremiumBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor('#334155'), # Cool Grey
            spaceAfter=10
        )

        story = []
        
        # Add Header/Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 10))
        
        # Add body content
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                story.append(Paragraph(line, body_style))
                story.append(Spacer(1, 5))
                
        # Optional: Add data tables for financial budgets / mining assays
        if table_data:
            table = Table(table_data, hAlign='LEFT')
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E293B')), # Dark slate
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 9.5),
                ('BOTTOMPADDING', (0,0), (-1,0), 6),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#F8FAFC')),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
                ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                ('FONTSIZE', (0,1), (-1,-1), 8.5),
                ('TOPPADDING', (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ]))
            story.append(Spacer(1, 15))
            story.append(table)

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def create_docx_from_content(self, title: str, content: str) -> bytes:
        """
        Generates a DOCX document using python-docx.
        """
        doc = Document()
        doc.add_heading(title, 0)
        
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line)
                
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def create_xlsx_from_dataframe(self, df: pd.DataFrame, sheet_name: str = "Sheet1") -> bytes:
        """
        Generates an Excel spreadsheet using pandas.
        """
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
        buffer.seek(0)
        return buffer.getvalue()

    def convert_csv_to_xlsx(self, csv_data: bytes) -> bytes:
        """
        Converts a CSV byte stream into an Excel (XLSX) byte stream.
        """
        csv_str = csv_data.decode("utf-8")
        df = pd.read_csv(io.StringIO(csv_str))
        return self.create_xlsx_from_dataframe(df)

    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extracts plain text from PDF bytes.
        """
        try:
            import pdfplumber
            text_runs = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_runs.append(page_text)
            return "\n".join(text_runs)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""

    def process_and_store_report(self, filename: str, content: str, file_type: str = "pdf", table_data: Optional[List[List[Any]]] = None) -> Dict[str, Any]:
        """
        Converts text content into requested report format, uploads to storage, and returns metadata.
        """
        file_type = file_type.lower()
        if file_type == "pdf":
            file_bytes = self.create_pdf_from_content(filename.replace(".pdf", ""), content, table_data)
            mime_type = "application/pdf"
        elif file_type == "docx":
            file_bytes = self.create_docx_from_content(filename.replace(".docx", ""), content)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_bytes = self.create_txt_from_content(filename, content)
            mime_type = "text/plain"

        # If minio_client is available, upload it
        storage_uri = f"local://reports/{filename}"
        if self.minio_client:
            storage_uri = self.minio_client.upload_file(filename, file_bytes, mime_type)

        return {
            "filename": filename,
            "storage_uri": storage_uri,
            "mime_type": mime_type,
            "size_bytes": len(file_bytes)
        }
